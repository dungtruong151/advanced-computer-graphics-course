from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Page setup: A4
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# --- Header ---
header_lines = [
    "INTERNATIONAL UNIVERSITY - VIETNAM NATIONAL UNIVERSITY HCMC",
    "SCHOOL OF COMPUTER SCIENCE AND ENGINEERING",
]
for line in header_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line)
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("\u2015" * 40)
run.font.size = Pt(11)
p.paragraph_format.space_after = Pt(4)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Advanced Computer Graphics (IT516)")
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
run.bold = True
p.paragraph_format.space_after = Pt(10)

# --- TITLE ---
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("PROJECT PROPOSAL")
run.bold = True
run.font.size = Pt(16)
run.font.name = 'Times New Roman'
run.font.color.rgb = RGBColor(0, 51, 102)
p.paragraph_format.space_after = Pt(4)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Normal Field Diffusion-Guided Hole Filling on 3D Meshes")
run.bold = True
run.font.size = Pt(14)
run.font.name = 'Times New Roman'
p.paragraph_format.space_after = Pt(14)

# --- Member ---
p = doc.add_paragraph()
run = p.add_run("Member:")
run.bold = True
run.font.size = Pt(12)
p.paragraph_format.space_after = Pt(1)

p = doc.add_paragraph()
run = p.add_run("    Truong Tri Dung \u2013 MITIU25208")
run.font.size = Pt(12)
p.paragraph_format.space_after = Pt(10)

# --- 1. Problem Statement ---
p = doc.add_paragraph()
run = p.add_run("1. Problem Statement")
run.bold = True
run.font.size = Pt(13)
p.paragraph_format.space_after = Pt(3)

desc = (
    "Hole filling is a fundamental problem in 3D mesh processing. "
    "Holes commonly arise during 3D scanning due to occlusion, sensor noise, "
    "or reflective surfaces. Existing methods such as Liepa (2003), advancing front, "
    "and volumetric diffusion typically produce flat or overly smoothed patches "
    "that do not conform naturally to the surrounding surface geometry."
)
p = doc.add_paragraph(desc)
p.paragraph_format.space_after = Pt(6)
for r in p.runs:
    r.font.size = Pt(12)

# --- 2. Proposed Method ---
p = doc.add_paragraph()
run = p.add_run("2. Proposed Method")
run.bold = True
run.font.size = Pt(13)
p.paragraph_format.space_after = Pt(3)

desc2 = (
    "This project proposes a novel method called Normal Field Diffusion-Guided "
    "Hole Filling (NFD-HF). The core idea is to propagate normal vectors from the "
    "hole boundary into the interior using heat diffusion on the mesh surface. "
    "The diffused normal field then guides vertex displacement to reconstruct "
    "a geometrically consistent surface patch. Unlike conventional approaches that "
    "rely solely on positional smoothing, NFD-HF leverages differential geometry "
    "information (normals and curvature) to predict the missing surface shape."
)
p = doc.add_paragraph(desc2)
p.paragraph_format.space_after = Pt(6)
for r in p.runs:
    r.font.size = Pt(12)

# --- 3. Pipeline ---
p = doc.add_paragraph()
run = p.add_run("3. Processing Pipeline")
run.bold = True
run.font.size = Pt(13)
p.paragraph_format.space_after = Pt(3)

steps = [
    ("Hole Detection",
     "identify boundary edges (edges belonging to exactly one triangle) "
     "and trace them into closed loops. Each loop represents one hole."),
    ("Initial Triangulation",
     "fill the hole with an initial mesh patch using ear clipping, "
     "then refine large triangles by inserting interior vertices."),
    ("Boundary Analysis",
     "compute vertex normals (area-weighted face normal averaging) and "
     "mean curvature (cotangent Laplacian) at each boundary vertex."),
    ("Normal Field Diffusion",
     "solve the heat equation \u2202n/\u2202t = \u03bb\u00b7\u0394n on the patch mesh "
     "using implicit Euler integration. Boundary normals are fixed as Dirichlet "
     "conditions; interior normals converge to a smooth interpolation."),
    ("Curvature-Guided Displacement",
     "displace each interior vertex along its diffused normal: "
     "v' = v + d\u00b7n, where d depends on the average boundary curvature "
     "and the geodesic distance to the boundary."),
    ("Local Smoothing",
     "apply constrained Laplacian smoothing (2\u20133 iterations) on the patch "
     "with boundary vertices held fixed, ensuring a smooth transition."),
]
for i, (title, detail) in enumerate(steps, 1):
    p = doc.add_paragraph()
    run = p.add_run(f"Step {i} \u2013 {title}: ")
    run.bold = True
    run.font.size = Pt(12)
    run2 = p.add_run(detail)
    run2.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)

# --- 4. Tools ---
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
run = p.add_run("4. Tools and Evaluation")
run.bold = True
run.font.size = Pt(13)
p.paragraph_format.space_after = Pt(3)

p = doc.add_paragraph(
    "Implementation: Python with NumPy, SciPy (sparse linear solver), "
    "and Trimesh. Results are exported as PLY files and visualized in MeshLab. "
    "Evaluation is performed on Stanford models (Bunny, Spot) with artificially "
    "created holes, using three metrics: Hausdorff Distance, RMS Error, "
    "and Normal Deviation (angular difference between filled and ground-truth normals)."
)
p.paragraph_format.space_after = Pt(0)
for r in p.runs:
    r.font.size = Pt(12)

# Save
output_path = r"c:\Users\ACE\Desktop\iu\Advanced Computer Graphics\NFD-HoleFill\NFD_HoleFill_Proposal_EN.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
