# -*- coding: utf-8 -*-
"""Generate NFD-HF Proposal v2 - same as v1, only replace Python with C++ plugin"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Copy exact style from v1
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

def add_normal(text):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    return p

# --- Exact same as v1 ---

# Title (fix typo "Diffusi+on" -> "Diffusion")
p = doc.add_paragraph('PROJECT PROPOSAL')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in p.runs:
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

p = doc.add_paragraph('Normal Field Diffusion-Guided Hole Filling on 3D Meshes')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in p.runs:
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

# Member
add_normal('Member:')
add_normal('\tTruong Tri Dung \u2013 MITIU25208')

# 1. Problem Statement (SAME as v1)
add_normal('1. Problem Statement')

add_normal(
    'Hole filling is a fundamental problem in 3D mesh processing. Holes commonly arise during '
    '3D scanning due to occlusion, sensor noise, or reflective surfaces. Existing methods such as '
    'Liepa (2003), advancing front, and volumetric diffusion typically produce flat or overly smoothed '
    'patches that do not conform naturally to the surrounding surface geometry.'
)

# 2. Proposed Method (SAME as v1)
add_normal('2. Proposed Method')

add_normal(
    'This project proposes a novel method called Normal Field Diffusion-Guided Hole Filling (NFD-HF). '
    'The core idea is to propagate normal vectors from the hole boundary into the interior using '
    'heat diffusion on the mesh surface. The diffused normal field then guides vertex displacement '
    'to reconstruct a geometrically consistent surface patch. Unlike conventional approaches that '
    'rely solely on positional smoothing, NFD-HF leverages differential geometry information '
    '(normals and curvature) to predict the missing surface shape.'
)

# 3. Processing Pipeline (SAME as v1)
add_normal('3. Processing Pipeline')

add_normal(
    'Step 1 \u2013 Hole Detection: identify boundary edges (edges belonging to exactly one triangle) '
    'and trace them into closed loops. Each loop represents one hole.'
)

add_normal(
    'Step 2 \u2013 Initial Triangulation: fill the hole with an initial mesh patch using ear clipping, '
    'then refine large triangles by inserting interior vertices.'
)

add_normal(
    'Step 3 \u2013 Boundary Analysis: compute vertex normals (area-weighted face normal averaging) '
    'and mean curvature (cotangent Laplacian) at each boundary vertex.'
)

add_normal(
    'Step 4 \u2013 Normal Field Diffusion: solve the heat equation \u2202n/\u2202t = \u03bb\u00b7\u0394n '
    'on the patch mesh using implicit Euler integration. Boundary normals are fixed as Dirichlet '
    'conditions; interior normals converge to a smooth interpolation.'
)

add_normal(
    'Step 5 \u2013 Curvature-Guided Displacement: displace each interior vertex along its diffused '
    "normal: v' = v + d\u00b7n, where d depends on the average boundary curvature and the geodesic "
    'distance to the boundary.'
)

add_normal(
    'Step 6 \u2013 Local Smoothing: apply constrained Laplacian smoothing (2\u20133 iterations) on '
    'the patch with boundary vertices held fixed, ensuring a smooth transition.'
)

# 4. Tools and Evaluation (CHANGED: Python -> C++ plugin)
add_normal('4. Tools and Evaluation')

add_normal(
    'Implementation: C++ MeshLab filter plugin using the VCG Library (mesh data structures, '
    'hole detection, normal/curvature computation, Laplacian smoothing) and Eigen (sparse matrix '
    'assembly and SimplicialLDLT solver for the heat diffusion linear system). The plugin is '
    'integrated into the MeshLab build tree via CMake and accessible through the Filters menu. '
    'Evaluation is performed on Stanford models (Bunny, Spot) with artificially created holes, '
    'using three metrics: Hausdorff Distance, RMS Error, and Normal Deviation (angular difference '
    'between filled and ground-truth normals).'
)

doc.save('NFD_HoleFill_Proposal_EN_v2.docx')
print('Saved NFD_HoleFill_Proposal_EN_v2.docx')
