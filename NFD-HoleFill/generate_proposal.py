# -*- coding: utf-8 -*-
"""Generate NFD-HF Proposal v2 (C++ Plugin version)"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# --- TITLE ---
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('PROJECT PROPOSAL')
run.bold = True
run.font.size = Pt(16)
run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Normal Field Diffusion-Guided Hole Filling on 3D Meshes')
run.bold = True
run.font.size = Pt(14)
run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('(MeshLab C++ Plugin Implementation)')
run.italic = True
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

# Member
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Member:')
run.bold = True
doc.add_paragraph('Truong Tri Dung \u2013 MITIU25208')

# --- 1. Problem Statement ---
doc.add_paragraph()
h = doc.add_heading('1. Problem Statement', level=1)
for run in h.runs:
    run.font.name = 'Times New Roman'

doc.add_paragraph(
    'Hole filling is a fundamental problem in 3D mesh processing. Holes commonly arise during '
    '3D scanning due to occlusion, sensor noise, or reflective surfaces. Existing methods such as '
    'Liepa (2003), advancing front, and volumetric diffusion typically produce flat or overly smoothed '
    'patches that do not conform naturally to the surrounding surface geometry.'
)

doc.add_paragraph(
    'Current hole-filling filters available in MeshLab (e.g., Close Holes) use simple triangulation '
    'strategies that ignore the differential geometry of the surrounding surface. This results in '
    'visible artifacts, especially on curved regions.'
)

# --- 2. Proposed Method ---
doc.add_paragraph()
h = doc.add_heading('2. Proposed Method', level=1)
for run in h.runs:
    run.font.name = 'Times New Roman'

doc.add_paragraph(
    'This project proposes a novel method called Normal Field Diffusion-Guided Hole Filling (NFD-HF), '
    'implemented as a native MeshLab C++ filter plugin using the VCG Library. '
    'The core idea is to propagate normal vectors from the hole boundary into the interior using '
    'heat diffusion on the mesh surface. The diffused normal field then guides vertex displacement '
    'to reconstruct a geometrically consistent surface patch.'
)

doc.add_paragraph(
    'Unlike conventional approaches that rely solely on positional smoothing, NFD-HF leverages '
    'differential geometry information (normals and curvature) to predict the missing surface shape. '
    'The method is integrated directly into MeshLab as a filter plugin, allowing users to apply it '
    'through the standard Filters menu alongside existing mesh processing tools.'
)

# --- 3. Processing Pipeline ---
doc.add_paragraph()
h = doc.add_heading('3. Processing Pipeline', level=1)
for run in h.runs:
    run.font.name = 'Times New Roman'

steps = [
    ('Step 1 \u2013 Hole Detection',
     'Identify boundary edges (edges belonging to exactly one triangle) and trace them into closed '
     'boundary loops using the VCG half-edge traversal. Each loop represents one hole. The plugin '
     'parameter max_hole_size allows users to filter holes by the number of boundary edges.'),
    ('Step 2 \u2013 Initial Triangulation',
     'Fill each detected hole with an initial mesh patch using the ear clipping algorithm '
     '(vcg::tri::Hole::EarCuttingFill). For large holes, additional interior vertices are inserted '
     'via edge midpoint refinement to ensure sufficient resolution for the subsequent diffusion step.'),
    ('Step 3 \u2013 Boundary Analysis',
     'Compute vertex normals at each boundary vertex using area-weighted face normal averaging '
     '(vcg::tri::UpdateNormal). Compute discrete mean curvature at boundary vertices using the '
     'cotangent Laplacian formula: H(v) = (1/2A) * sum[(cot alpha_ij + cot beta_ij) * (v_i - v_j)].'),
    ('Step 4 \u2013 Normal Field Diffusion (Core Novelty)',
     'Solve the heat equation on the patch mesh to diffuse boundary normals into the interior. '
     'The system (I - lambda * dt * L) * n^(t+1) = n^(t) is solved using the Eigen sparse linear '
     'solver (SimplicialLDLT), where L is the cotangent Laplacian matrix of the patch. Boundary '
     'normals are fixed as Dirichlet conditions. The diffusion coefficient lambda and number of '
     'iterations are exposed as plugin parameters. After each iteration, normals are re-normalized.'),
    ('Step 5 \u2013 Curvature-Guided Displacement',
     'Displace each interior vertex along its diffused normal: v_new = v + d * n_diffused, where '
     'd = H_avg * g(dist). H_avg is the average mean curvature at the boundary, dist is the '
     'normalized topological distance from the vertex to the boundary (0 at boundary, 1 at center), '
     'and g(t) = t * (1-t) is a weighting function that peaks at the hole center.'),
    ('Step 6 \u2013 Local Smoothing',
     'Apply 2\u20133 iterations of constrained Laplacian smoothing (vcg::tri::Smooth) on the patch '
     'with boundary vertices held fixed, ensuring a smooth C0/C1 transition between the filled '
     'region and the original mesh surface.'),
]

for title, desc in steps:
    p = doc.add_paragraph()
    run = p.add_run(title + ': ')
    run.bold = True
    run.font.name = 'Times New Roman'
    run = p.add_run(desc)
    run.font.name = 'Times New Roman'

# --- 4. Implementation Details ---
doc.add_paragraph()
h = doc.add_heading('4. Implementation Details', level=1)
for run in h.runs:
    run.font.name = 'Times New Roman'

p = doc.add_paragraph()
run = p.add_run('Platform: ')
run.bold = True
p.add_run('MeshLab C++ filter plugin')

p = doc.add_paragraph()
run = p.add_run('Build System: ')
run.bold = True
p.add_run('CMake (integrated into MeshLab build tree)')

p = doc.add_paragraph()
run = p.add_run('Core Libraries:')
run.bold = True

libs = [
    ('VCG Library',
     'mesh data structures, half-edge traversal, hole detection (vcg::tri::Hole), '
     'normal/curvature computation, Laplacian smoothing'),
    ('Eigen',
     'sparse matrix assembly, SimplicialLDLT solver for the heat diffusion linear system'),
]

for lib_name, lib_desc in libs:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(lib_name + ': ')
    run.bold = True
    run.font.name = 'Times New Roman'
    run = p.add_run(lib_desc)
    run.font.name = 'Times New Roman'

# Plugin parameters table
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Plugin Parameters (exposed in MeshLab UI):')
run.bold = True

table = doc.add_table(rows=5, cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['Parameter', 'Type', 'Default', 'Description']
for i, h_text in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h_text
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)

params = [
    ['max_hole_size', 'int', '100', 'Maximum boundary edges per hole'],
    ['diffusion_lambda', 'float', '0.5', 'Diffusion coefficient'],
    ['diffusion_iters', 'int', '50', 'Number of diffusion iterations'],
    ['smooth_iters', 'int', '3', 'Laplacian smoothing iterations'],
]

for ri, row_data in enumerate(params):
    for ci, val in enumerate(row_data):
        cell = table.rows[ri + 1].cells[ci]
        cell.text = val
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)

# File structure
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Plugin File Structure:')
run.bold = True

code_text = (
    'meshlab/src/meshlabplugins/filter_nfd_holefill/\n'
    '  CMakeLists.txt\n'
    '  filter_nfd_holefill.h        // Plugin class declaration\n'
    '  filter_nfd_holefill.cpp      // Plugin registration, UI parameters\n'
    '  nfd_algorithm.h              // NFD-HF algorithm implementation\n'
)
p = doc.add_paragraph()
run = p.add_run(code_text)
run.font.name = 'Consolas'
run.font.size = Pt(10)

# --- 5. Evaluation ---
doc.add_paragraph()
h = doc.add_heading('5. Evaluation', level=1)
for run in h.runs:
    run.font.name = 'Times New Roman'

p = doc.add_paragraph()
run = p.add_run('Test Models: ')
run.bold = True
p.add_run(
    'Stanford Bunny (34,834 vertices) and Spot (3,225 vertices). '
    'Holes are artificially created by removing clusters of faces, preserving the original mesh '
    'as ground truth for quantitative comparison.'
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Metrics:')
run.bold = True

metrics_table = doc.add_table(rows=4, cols=2)
metrics_table.style = 'Table Grid'
metrics_table.alignment = WD_TABLE_ALIGNMENT.CENTER

m_headers = ['Metric', 'Description']
for i, h_text in enumerate(m_headers):
    cell = metrics_table.rows[0].cells[i]
    cell.text = h_text
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)

metrics_data = [
    ['Hausdorff Distance', 'Maximum geometric deviation between filled surface and ground truth'],
    ['RMS Error', 'Root mean square of vertex-to-surface distances'],
    ['Normal Deviation', 'Mean angular difference between filled and ground-truth face normals'],
]

for ri, row_data in enumerate(metrics_data):
    for ci, val in enumerate(row_data):
        cell = metrics_table.rows[ri + 1].cells[ci]
        cell.text = val
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Comparison Baseline: ')
run.bold = True
p.add_run(
    'MeshLab built-in Close Holes filter (flat triangulation + Laplacian smoothing). '
    'Both Hausdorff distance computation and visual comparison are performed directly within MeshLab '
    'using the Hausdorff Distance filter (Filters > Sampling > Hausdorff Distance).'
)

# --- 6. Expected Results ---
doc.add_paragraph()
h = doc.add_heading('6. Expected Results', level=1)
for run in h.runs:
    run.font.name = 'Times New Roman'

doc.add_paragraph(
    'The NFD-HF plugin is expected to produce filled surfaces that better conform to the surrounding '
    'geometry compared to the built-in Close Holes filter, particularly on curved regions such as '
    'the bunny ears and body. The normal field diffusion should yield lower Hausdorff distance, '
    'lower RMS error, and smaller normal deviation compared to the baseline.'
)

# --- 7. Timeline ---
doc.add_paragraph()
h = doc.add_heading('7. Timeline', level=1)
for run in h.runs:
    run.font.name = 'Times New Roman'

timeline_table = doc.add_table(rows=6, cols=2)
timeline_table.style = 'Table Grid'
timeline_table.alignment = WD_TABLE_ALIGNMENT.CENTER

t_headers = ['Week', 'Task']
for i, h_text in enumerate(t_headers):
    cell = timeline_table.rows[0].cells[i]
    cell.text = h_text
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)

timeline_data = [
    ['1\u20132', 'Set up MeshLab build environment, compile from source, create plugin skeleton'],
    ['3\u20134', 'Implement Steps 1\u20133 (hole detection, triangulation, boundary analysis)'],
    ['5\u20136', 'Implement Step 4 (normal field diffusion with Eigen solver) \u2013 core novelty'],
    ['7\u20138', 'Implement Steps 5\u20136 (displacement + smoothing), testing and evaluation'],
    ['9\u201310', 'Final report, presentation preparation, demo in MeshLab'],
]

for ri, row_data in enumerate(timeline_data):
    for ci, val in enumerate(row_data):
        cell = timeline_table.rows[ri + 1].cells[ci]
        cell.text = val
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)

# --- 8. References ---
doc.add_paragraph()
h = doc.add_heading('8. References', level=1)
for run in h.runs:
    run.font.name = 'Times New Roman'

refs = [
    '[1] P. Liepa, "Filling holes in meshes," in Proc. Eurographics Symposium on Geometry Processing, 2003.',
    '[2] P. Cignoni et al., "MeshLab: an open-source mesh processing tool," in Proc. Eurographics Italian Chapter Conference, 2008.',
    '[3] The VCG Library, Visual Computing Lab, ISTI-CNR. https://github.com/cnr-isti-vclab/vcglib',
    '[4] K. Crane, C. Weischedel, M. Wardetzky, "The heat method for distance computation," Communications of the ACM, 2017.',
    '[5] G. Guennebaud, B. Jacob et al., "Eigen v3," http://eigen.tuxfamily.org, 2010.',
]

for ref in refs:
    p = doc.add_paragraph(ref)
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)

doc.save('NFD_HoleFill_Proposal_EN_v2.docx')
print('Saved NFD_HoleFill_Proposal_EN_v2.docx')
